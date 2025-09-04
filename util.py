from docx import Document
from indico.modules.events.models.events import Event
from indico.modules.events.papers.models.revisions import PaperRevisionState
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from collections import defaultdict
from typing import List, Dict, Optional, Tuple
from datetime import date, datetime


class DocxGenerator:
    """Базовый класс для генерации DOCX документов"""
    
    # Константы для настроек документа
    MARGINS = {
        'left': Inches(0.79),
        'right': Inches(0.39),
        'top': Inches(0.79),
        'bottom': Inches(0.79)
    }
    
    FONT_SETTINGS = {
        'name': 'Times New Roman',
        'size': Pt(14),
        'color': RGBColor(0, 0, 0)
    }
    
    LINE_SPACING = 1.5
    
    # Перевод месяцев
    MONTH_TRANSLATIONS = {
        'January': 'января', 'February': 'февраля', 'March': 'марта', 'April': 'апреля',
        'May': 'мая', 'June': 'июня', 'July': 'июля', 'August': 'августа',
        'September': 'сентября', 'October': 'октября', 'November': 'ноября', 'December': 'декабря'
    }
    
    def __init__(self, event_id: int):
        self.event = Event.get(event_id)
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self) -> None:
        """Настройка базовых параметров документа"""
        for section in self.doc.sections:
            section.left_margin = self.MARGINS['left']
            section.right_margin = self.MARGINS['right']
            section.top_margin = self.MARGINS['top']
            section.bottom_margin = self.MARGINS['bottom']
    
    def _set_black_color(self, element) -> None:
        """Установка черного цвета для всех элементов"""
        if hasattr(element, 'runs'):
            for run in element.runs:
                run.font.color.rgb = self.FONT_SETTINGS['color']
        if hasattr(element, 'paragraphs'):
            for paragraph in element.paragraphs:
                self._set_black_color(paragraph)
        if hasattr(element, 'tables'):
            for table in element.tables:
                self._set_black_color(table)
    
    def _format_russian_date(self, date_obj: date, include_time: bool = False) -> str:
        """Форматирование даты на русском языке"""
        if include_time:
            date_str = date_obj.strftime('%d %B %Y г., %H-%M')
        else:
            date_str = date_obj.strftime('%d %B %Y г.')
        
        for eng, rus in self.MONTH_TRANSLATIONS.items():
            date_str = date_str.replace(eng, rus)
        
        return date_str
    
    def _get_contributions_by_date(self) -> Tuple[Dict[date, List], List]:
        """Группировка докладов по дате и отдельно без времени"""
        contributions_with_time = [c for c in self.event.contributions 
                                 if not c.is_deleted and c.start_dt]
        contributions_without_time = [c for c in self.event.contributions 
                                    if not c.is_deleted and not c.start_dt]
        
        date_groups = defaultdict(list)
        for contrib in contributions_with_time:
            date_key = contrib.start_dt.date()
            date_groups[date_key].append(contrib)
        
        # Сортировка докладов внутри каждой даты
        for date_key in date_groups:
            date_groups[date_key].sort(key=lambda x: x.start_dt)
        
        return dict(sorted(date_groups.items())), contributions_without_time
    
    def _get_speaker_name(self, person) -> str:
        """Форматирование имени докладчика"""
        middle_initial = f".{person.first_name[1]}" if len(person.first_name) > 1 else ""
        return f"{person.last_name} {person.first_name[0]}{middle_initial}"
    
    def _get_full_name(self, person) -> str:
        """Полное имя с отчеством если есть"""
        if hasattr(person, 'middle_name') and person.middle_name:
            return f"{person.last_name} {person.first_name} {person.middle_name}"
        return f"{person.first_name} {person.last_name}"
    
    def _determine_student_status(self, person) -> str:
        """Определение статуса участника"""
        if not person.affiliation:
            return 'Не указан'
        
        affiliation_lower = person.affiliation.lower()
        
        student_keywords = ['студент', 'student', 'бакалавр', 'bachelor', '1 курс', '2 курс', '3 курс', '4 курс']
        master_keywords = ['магистр', 'master', 'магистрант', '5 курс', '6 курс']
        
        for keyword in student_keywords:
            if keyword in affiliation_lower:
                return 'Студент'
        
        for keyword in master_keywords:
            if keyword in affiliation_lower:
                return 'Магистр'
        
        return person.affiliation
    
    def _add_heading(self, text: str, level: int = 0, alignment: int = WD_ALIGN_PARAGRAPH.CENTER) -> None:
        """Добавление заголовка"""
        heading = self.doc.add_heading(text, level)
        heading.alignment = alignment
        for run in heading.runs:
            run.font.color.rgb = self.FONT_SETTINGS['color']
    
    def _add_centered_paragraph(self, text: str, bold: bool = False) -> None:
        """Добавление центрированного параграфа"""
        paragraph = self.doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold and paragraph.runs:
            paragraph.runs[0].font.bold = True
            paragraph.runs[0].font.color.rgb = self.FONT_SETTINGS['color']
    
    def _apply_document_styling(self) -> None:
        """Применение стилей ко всему документу"""
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.FONT_SETTINGS['name']
                run.font.size = self.FONT_SETTINGS['size']
                run.font.color.rgb = self.FONT_SETTINGS['color']
            paragraph.paragraph_format.line_spacing = self.LINE_SPACING
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.FONT_SETTINGS['name']
                            run.font.size = self.FONT_SETTINGS['size']
                            run.font.color.rgb = self.FONT_SETTINGS['color']
                        paragraph.paragraph_format.line_spacing = self.LINE_SPACING
        
        self._set_black_color(self.doc)
    
    def _save_to_bytes(self) -> bytes:
        """Сохранение документа в bytes"""
        f = BytesIO()
        self.doc.save(f)
        return f.getvalue()


class ContributionsListGenerator(DocxGenerator):
    """Генератор списка докладов"""
    
    def generate(self) -> bytes:
        """Генерация документа со списком докладов"""
        self._add_heading('СПИСОК ДОКЛАДОВ', 0)
        self._add_centered_paragraph(f'"{self.event.title}"', bold=True)
        self.doc.add_paragraph()
        
        date_groups, no_time_contribs = self._get_contributions_by_date()
        
        # Доклады с временем
        if date_groups:
            self._add_date_grouped_contributions(date_groups)
        
        # Доклады без времени
        if no_time_contribs:
            self._add_no_time_contributions(no_time_contribs)
        
        self._apply_document_styling()
        return self._save_to_bytes()
    
    def _add_date_grouped_contributions(self, date_groups: Dict[date, List]) -> None:
        """Добавление докладов сгруппированных по дате"""
        sorted_dates = sorted(date_groups.keys())
        
        for i, date_key in enumerate(sorted_dates, 1):
            date_contributions = date_groups[date_key]
            date_str = self._format_russian_date(date_key)
            
            meeting_title = f'Заседание {i}' if len(sorted_dates) > 1 else 'Заседание'
            self._add_heading(meeting_title, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            date_para = self.doc.add_paragraph(date_str)
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            self._create_contributions_table(date_contributions)
            self.doc.add_paragraph()
    
    def _add_no_time_contributions(self, contributions: List) -> None:
        """Добавление докладов без указанного времени"""
        self._add_heading('Доклады без указанного времени', level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        self._create_contributions_table(contributions)
    
    def _create_contributions_table(self, contributions: List) -> None:
        """Создание таблицы с докладами"""
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        headers = ['№', 'Фамилия и инициалы докладчика, название доклада', 
                  'Статус (магистр / студент)', 'Решение']
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Заполнение таблицы
        row_number = 1
        for contribution in sorted(contributions, key=lambda x: x.title.lower() if x.title else ''):
            speakers = [link.person for link in contribution.person_links if link.is_speaker]
            
            if not speakers:
                continue
                
            for speaker in speakers:
                row = table.add_row().cells
                
                # Номер
                row[0].text = str(row_number)
                row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Докладчик и название
                speaker_name = self._get_speaker_name(speaker)
                contribution_title = contribution.title or 'Без названия'
                row[1].text = f"{speaker_name}. {contribution_title}"
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Статус
                status = self._determine_student_status(speaker)
                row[2].text = status
                row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Решение (пустое)
                row[3].text = ''
                row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                row_number += 1


class ConferenceReportGenerator(DocxGenerator):
    """Генератор отчета о конференции"""
    
    def generate(self) -> bytes:
        """Генерация отчета о конференции"""
        self._add_heading('ОТЧЕТ О ПРОВЕДЕНИИ КОНФЕРЕНЦИИ', 0)
        self._add_centered_paragraph(f'"{self.event.title}"', bold=True)
        self.doc.add_paragraph()
        
        date_groups, no_time_contribs = self._get_contributions_by_date()
        
        # Доклады с временем
        if date_groups:
            self._add_date_grouped_contributions(date_groups)
        
        # Доклады без времени
        if no_time_contribs:
            self._add_no_time_contributions(no_time_contribs)
        
        self._apply_document_styling()
        return self._save_to_bytes()
    
    def _add_date_grouped_contributions(self, date_groups: Dict[date, List]) -> None:
        """Добавление докладов сгруппированных по дате"""
        sorted_dates = sorted(date_groups.keys())
        
        for i, date_key in enumerate(sorted_dates, 1):
            date_contributions = date_groups[date_key]
            date_str = self._format_russian_date(date_key, include_time=True)
            
            meeting_title = f'Заседание {i}' if len(sorted_dates) > 1 else 'Заседание'
            self._add_heading(meeting_title, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            date_para = self.doc.add_paragraph(date_str)
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            self._add_contributions_list(date_contributions)
            self.doc.add_paragraph()
    
    def _add_no_time_contributions(self, contributions: List) -> None:
        """Добавление докладов без указанного времени"""
        self._add_heading('Доклады без указанного времени', level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        self._add_contributions_list(contributions)
    
    def _add_contributions_list(self, contributions: List) -> None:
        """Добавление списка докладов в виде параграфов"""
        row_number = 1
        for contribution in sorted(contributions, key=lambda x: x.title.lower() if x.title else ''):
            speakers = [link.person for link in contribution.person_links if link.is_speaker]
            
            if not speakers:
                continue
                
            for speaker in speakers:
                p = self.doc.add_paragraph()
                
                run_number = p.add_run(f"{row_number}. ")
                run_number.bold = True
                
                speaker_name = self._get_speaker_name(speaker)
                run_name = p.add_run(speaker_name)
                run_name.bold = True
                
                contribution_title = contribution.title or 'Без названия'
                p.add_run(f". {contribution_title}")
                
                row_number += 1


class PublicationsListGenerator(DocxGenerator):
    """Генератор списка публикаций"""
    
    def generate(self) -> bytes:
        """Генерация списка публикаций"""
        self._add_heading('СПИСОК ПУБЛИКАЦИЙ', 0)
        self._add_centered_paragraph(f'"{self.event.title}"', bold=True)
        self.doc.add_paragraph()
        
        date_groups, no_time_contribs = self._get_contributions_by_date()
        has_publications = False
        
        # Публикации с временем
        if date_groups:
            has_publications = self._add_date_grouped_publications(date_groups) or has_publications
        
        # Публикации без времени
        if no_time_contribs:
            has_publications = self._add_no_time_publications(no_time_contribs) or has_publications
        
        if not has_publications:
            p = self.doc.add_paragraph()
            p.add_run("Статьи, принятые к публикации, не найдены.")
            p.italic = True
        
        self._apply_document_styling()
        return self._save_to_bytes()
    
    def _add_date_grouped_publications(self, date_groups: Dict[date, List]) -> bool:
        """Добавление публикаций сгруппированных по дате"""
        sorted_dates = sorted(date_groups.keys())
        has_publications = False
        
        for i, date_key in enumerate(sorted_dates, 1):
            date_contributions = date_groups[date_key]
            date_str = self._format_russian_date(date_key)
            
            meeting_title = f'Заседание {i}.' if len(sorted_dates) > 1 else 'Заседание.'
            self._add_heading(meeting_title, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            date_para = self.doc.add_paragraph(date_str)
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            has_date_publications = self._add_publications_list(date_contributions)
            has_publications = has_publications or has_date_publications
            self.doc.add_paragraph()
        
        return has_publications
    
    def _add_no_time_publications(self, contributions: List) -> bool:
        """Добавление публикаций без указанного времени"""
        self._add_heading('Доклады без указанного времени', level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        return self._add_publications_list(contributions)
    
    def _add_publications_list(self, contributions: List) -> bool:
        """Добавление списка публикаций"""
        row_number = 1
        has_publications = False
        
        for contribution in sorted(contributions, key=lambda x: x.title.lower() if x.title else ''):
            # Проверяем, есть ли принятая статья
            if (hasattr(contribution, '_accepted_paper_revision') and 
                contribution._accepted_paper_revision and
                hasattr(contribution._accepted_paper_revision, 'state') and
                contribution._accepted_paper_revision.state == PaperRevisionState.accepted):
                
                authors = [link.person for link in contribution.person_links if link.is_speaker]
                
                if not authors:
                    continue
                
                for author in authors:
                    p = self.doc.add_paragraph()
                    
                    run_number = p.add_run(f"    {row_number}. ")
                    run_number.bold = True
                    
                    full_name = self._get_full_name(author)
                    run_name = p.add_run(full_name)
                    run_name.bold = True
                    
                    # Добавляем affiliation
                    if author.affiliation:
                        p.add_run(f", {author.affiliation}")
                    
                    p.add_run("\n")
                    
                    article_title = contribution.title or 'Без названия'
                    p.add_run(article_title)
                    
                    row_number += 1
                    has_publications = True
        
        return has_publications


# Функции-обертки для обратной совместимости
def generate_docx_list(event_id: int) -> bytes:
    """Генерация списка докладов"""
    generator = ContributionsListGenerator(event_id)
    return generator.generate()

def generate_docx_report(event_id: int) -> bytes:
    """Генерация отчета о конференции"""
    generator = ConferenceReportGenerator(event_id)
    return generator.generate()

def generate_docx_papers(event_id: int) -> bytes:
    """Генерация списка публикаций"""
    generator = PublicationsListGenerator(event_id)
    return generator.generate()